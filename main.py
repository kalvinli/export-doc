from flask import Flask, jsonify, request, render_template, send_file
import os
import platform
import shutil
import subprocess
import json
import time
import uuid
import base64
import re
from requests_toolbelt import MultipartEncoder
from docx import Document
from docx.shared import Pt, Cm
from docx.shared import RGBColor
from docx.oxml.ns import nsmap, qn
from docx.oxml import OxmlElement

# from docx2pdf import convert
from apscheduler.schedulers.background import BackgroundScheduler
from base_class.base_api import BaseClass
from base_class.generator import generate_qrcode, generate_barcode


app = Flask(__name__, static_folder="static", static_url_path="/static")
app.config['JSON_AS_ASCII'] = False
app.config['JSON_SORT_KEYS'] = False
app.json.ensure_ascii = False

# 当前脚本的目录
fp = os.path.dirname(os.path.abspath(__file__))

## 定义模板文件的保存路径和文件名尾缀
UPLOAD_FOLDER = os.path.join(fp, 'template_files')
ALLOWED_EXTENSIONS = {'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

## 独立生成二维码和条形码的保存路径
GENERATE_FOLDER = os.path.join(fp, 'generate_files')
app.config['GENERATE_FOLDER'] = GENERATE_FOLDER

# timestamp = time.time()
# local_time = time.localtime()
# formatted_local_time = time.strftime('%Y-%m-%d %H:%M:%S', local_time)
# formatted_local_time = time.strftime('%Y-%m-%d', local_time)



## 基于上传的模板将多维表格中的记录数据导出到word文件，并回传到当前记录的附件字段中
def export_to_doc(app_token, personal_base_token, table_id, record_id, info_json,
                  file_name, file_field, field_id_map, file_type, template_file_type):
    '''
        基于上传的模板将多维表格中的记录数据导出到word文件，并回传到当前记录的附件字段中\r\n
        pramas:\r\n
        - app_token: 多维表格ID\r\n
        - personal_base_token: 多维表格授权码\r\n
        - table_id: 数据表ID\r\n
        - record_id: 记录ID\r\n
        - info_json: 字段名与字段值的映射\r\n
        - file_name: 导出成附件的文件名，默认设定为多维表格中某个字段中的值\r\n
        - file_field: 导出成附件后回传的附件字段名\r\n
        - field_id_map: 字段名与字段ID的映射\r\n
        - file_type: 导出文件类型，默认为docx，当前在linux环境下面部署后导出为pdf有问题\r\n
    '''

    msg = "生成附件成功"

    
    # 配置命名空间
    nsmap.update({
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "image": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wpc": "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "v": "urn:schemas-microsoft-com:vml",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        # "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
        # "o": "urn:schemas-microsoft-com:office:office",
        # "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        # "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
        # "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
        # "w10": "urn:schemas-microsoft-com:office:word",
        # "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
        # "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
        # "wpi": "http://schemas.microsoft.com/office/word/2010/wordprocessingInk",
        # "wne": "http://schemas.microsoft.com/office/word/2006/wordml",
        # "wpsCustomData": "http://www.wps.cn/officeDocument/2013/wpsCustomData"
    })

    # print(info_json)
    # print("*"*30)
    # print(field_id_map)
    # print("*"*30)

    for key, value in field_id_map.items():
        if key not in info_json:
            info_json[key] = ""

    # print(info_json)
    # print("*"*30)

    # 每个多维表格主路径
    main_path = os.path.join(app.config['UPLOAD_FOLDER'], personal_base_token)
    # print(str(main_path))
    # print("*"*30)

    # 每个多维表格数据表主路径
    table_main_path = os.path.join(main_path, table_id)

    # 模板文件路径
    template_file_path = os.path.join(table_main_path, template_file_type, "template.docx")
    # print(template_file_path)

    # 临时生成的主目录，以file_name为文件夹名
    personal_main_path = os.path.join(table_main_path, file_name)
    # print(personal_main_path)

    # 临时生成的word文件路径
    target_file_path = os.path.join(personal_main_path, file_name + ".docx")
    # print(target_file_path)

    # 个人图片文件路径
    image_file_path = os.path.join(personal_main_path, file_name + ".jpg")

    # 印章图片文件路径
    seal_image_file_path = os.path.join(personal_main_path, file_name + ".png")

    # 如果模板文件不存在，则直接返回
    if not os.path.isfile(template_file_path):
        print("模板文件不存在，请先上传模板文件")
        return "模板文件不存在，请先上传模板文件"

    try:
        os.mkdir(personal_main_path)
    except Exception as e:
        print(e)
        print("当前文件夹已存在，删除文件夹下面的所有文件")
        delete_files_in_directory(personal_main_path)

    # 从模板文件创建一个副本文件
    shutil.copy(template_file_path, target_file_path)


    # 基于副本文件初始化一个文档实例
    doc = Document(target_file_path)

    # 查找所有可能包含文本框的XML元素
    search_paths = [
        './/w:txbxContent//w:t',  # 新版Word文本框
        './/v:textbox//w:t',  # 旧版Word文本框
        './/wps:txbx//w:t',  # Word 2010+ 文本框
        './/wpc:txbx//w:t'  # 绘图画布中的文本框
    ]

    ## 遍历文档中的所有文本段落
    for paragraph in doc.paragraphs:
        index = 0
        # 根据获取段落文本，遍历字段名列表进行占位字符的替换
        for key, value in info_json.items():
            # print(key, value)

            # 判断段落文本中是否包含有`{{字段名` 这样的信息，如果存在，则表示存在占位符否则不处理当前段落文本
            if '{{' + key in paragraph.text:
                # print(key, paragraph.text)

                # 遍历段落中的所有文本片断
                for run in paragraph.runs:

                    # 判断文本片断中是否包含有`{{字段名` 这样的信息，如果存在，则表示存在占位符
                    if '{{' + key in run.text:
                        # print(key, run.text)

                        # 获取当前文本片断的样式，当前只处理字体大小、颜色、加粗和斜体四种样式
                        font_size = run.font.size  # 假设所有格式相同，这里仅取第一个run的格式
                        color = run.font.color.rgb  # 保存颜色，如果有的话
                        bold = run.font.bold is not None  # 保存粗体状态
                        italic = run.font.italic is not None  # 保存斜体状态

                        # 根据占位符格式对文本片断进行替换
                        try:
                            # 如果占位符包含有`:image`，替换占位符为图片
                            if ":image" in run.text:
                                # 将文本片断中的`{{`和`}}`替换为空，保留有用信息
                                run_text = run.text.replace("{{", "").replace(
                                    "}}", "")
                                # 将以上信息用`:`分割，生成列表
                                key_split = run_text.split(":")
                                key = key_split[0]
                                # 如果列表的长度为3，则以上信息中包含有图片尺寸
                                if len(key_split) == 3:
                                    # 对图片尺寸进行分割后并进行变量赋值
                                    size = key_split[2].split("*")
                                    width = float(size[0])
                                    height = float(size[1])

                                # 如果列表的长度不为3，则以上信息中不包含有图片尺寸，则不处理图片大小
                                else:
                                    width = None
                                    height = None

                                # 生成多维表格附件下载的extra信息，并进行附件下载，返回附件文件的二进制流信息
                                for value_item in value:
                                    file_token = value_item.get("file_token")
                                    name = value_item.get("name")
                                    extra = {
                                        "bitablePerm": {
                                            "tableId": table_id,
                                            "attachments": {
                                                field_id_map[key]: {
                                                    record_id: [file_token]
                                                }
                                            }
                                        }
                                    }
                                    attachment_resp = BaseClass(
                                    ).download_attachment(personal_base_token, file_token,
                                                        extra)
                                    # print(attachment_resp)

                                    # 将二进制流信息写入到个人生成的图片附件中
                                    image_file_path = os.path.join(personal_main_path, name)
                                    with open(image_file_path, 'wb') as file:
                                        file.write(attachment_resp.content)
                                    file.close()

                                    # 将图片替换到图片占位符所在位置，并把原有的占位符文本置为空
                                    try:
                                        paragraph.add_run().add_picture(
                                            image_file_path,
                                            width=Cm(width),
                                            height=Cm(height))
                                    except Exception as e:
                                        paragraph.add_run().add_picture(
                                            image_file_path)
                                        
                                    os.remove(image_file_path)

                                run.text = ""

                            # 如果占位符包含有`:barcode`，替换占位符为条形码
                            elif ":barcode" in run.text:
                                run_text = run.text.replace("{{", "").replace(
                                    "}}", "")
                                key_split = run_text.split(":")
                                key = key_split[0]
                                if len(key_split) == 3:
                                    size = key_split[2].split("*")
                                    width = float(size[0])
                                    height = float(size[1])
                                else:
                                    width = None
                                    height = None

                                # 调用接口生成条形码，默认为模板文件相同路径，并以当前字段的值作为文件名，默认使用`code128`的样式，并重新设定条形码的文件路径
                                barcode_file_name = generate_barcode(
                                    value, 'code128', None, personal_main_path)
                                barcode_file_path = os.path.join(
                                    personal_main_path, barcode_file_name)

                                # 将条形码替换到条形码占位符所在位置，并把原有的占位符文本置为空
                                try:
                                    paragraph.add_run().add_picture(
                                        barcode_file_path,
                                        width=Cm(width),
                                        height=Cm(height))
                                except Exception as e:
                                    paragraph.add_run().add_picture(
                                        barcode_file_path)
                                run.text = ""

                            # 如果占位符包含有`:qrcode`，替换占位符为二维码
                            elif ":qrcode" in run.text:
                                run_text = run.text.replace("{{", "").replace(
                                    "}}", "")
                                key_split = run_text.split(":")
                                key = key_split[0]
                                if len(key_split) == 3:
                                    size = key_split[2].split("*")
                                    width = float(size[0])
                                    height = float(size[1])
                                else:
                                    width = None
                                    height = None

                                # 调用接口生成二维码，默认为模板文件相同路径，并以当前字段的值作为文件名，默认使用`code128`的样式，并重新设定条形码的文件路径
                                qrcode_file_name = generate_qrcode(
                                    value, {}, personal_main_path)
                                qrcode_file_path = os.path.join(
                                    personal_main_path, qrcode_file_name)

                                # 将二维码替换到二维码占位符所在位置，并把原有的占位符文本置为空
                                try:
                                    paragraph.add_run().add_picture(
                                        qrcode_file_path,
                                        width=Cm(width),
                                        height=Cm(height))
                                except Exception as e:
                                    paragraph.add_run().add_picture(
                                        qrcode_file_path)
                                run.text = ""

                            # 如果不是以上三种情况，则直接替换为对应字段的值
                            else:
                                run.text = run.text.replace(
                                    "{{" + key + "}}", value, 1)

                        # 如果替换失败，则将当前文本片断置为空，继续后面的执行
                        except Exception as e:
                            run.text = ""
                            print(e)

                        # 应用之前保存的样式
                        if font_size:  # 如果存在字体大小，设置字体大小
                            run.font.size = Pt(font_size.pt)
                        if color:  # 如果存在颜色，则设置颜色
                            run.font.color.rgb = RGBColor(*color)
                        if bold:  # 如果存在粗体，则设置粗体
                            run.font.bold = bold
                        if italic:  # 如果存在斜体，则设置斜体
                            run.font.italic = italic
                        break  # 只考虑第一个出现的占位符

                index = index + 1

            # 判断是否包含浮动文本框，并进行相应的替换操作
            else:
                # 使用正确的命名空间查询浮动文本框
                para_element = paragraph._element
                flag = False
                for path in search_paths:
                    elements = para_element.xpath(path)
                    for element in elements:
                        try:
                            if '{{' + key in element.text:
                                # print(key, element.text)
                                flag = True

                                # 如果占位符包含有`:image`，替换占位符为图片
                                if ":image" in element.text:
                                    # 将文本片断中的`{{`和`}}`替换为空，保留有用信息
                                    element_text = element.text.replace(
                                        "{{", "").replace("}}", "")
                                    # 将以上信息用`:`分割，生成列表
                                    key_split = element_text.split(":")
                                    key = key_split[0]
                                    # 如果列表的长度为3，则以上信息中包含有图片尺寸
                                    if len(key_split) == 3:
                                        # 对图片尺寸进行分割后并进行变量赋值
                                        size = key_split[2].split("*")
                                        width = float(size[0])
                                        height = float(size[1])

                                    # 如果列表的长度不为3，则以上信息中不包含有图片尺寸，则不处理图片大小
                                    else:
                                        width = None
                                        height = None

                                    # 生成多维表格附件下载的extra信息，并进行附件下载，返回附件文件的二进制流信息
                                    for value_item in value:
                                        file_token = value_item.get("file_token")
                                        name = value_item.get("name")
                                        extra = {
                                            "bitablePerm": {
                                                "tableId": table_id,
                                                "attachments": {
                                                    field_id_map[key]: {
                                                        record_id: [file_token]
                                                    }
                                                }
                                            }
                                        }
                                        attachment_resp = BaseClass(
                                        ).download_attachment(
                                            personal_base_token, file_token, extra)
                                        # print(attachment_resp)

                                        # 将二进制流信息写入到生成印章的图片附件中
                                        image_file_path = os.path.join(personal_main_path, name)
                                        with open(image_file_path,
                                                'wb') as file:
                                            file.write(attachment_resp.content)
                                        file.close()

                                        try:
                                            # 清空原有文本
                                            parent = element.getparent()
                                            # parent.remove(element)

                                            # 添加图片关系
                                            image_part = doc.part.package.get_or_add_image_part(image_file_path)
                                            r_id = doc.part.relate_to(image_part, nsmap['image'])
                                            
                                            # 创建图片元素
                                            image_element = create_image_element(r_id, width, height)
                                            
                                            # 构建标准文档结构
                                            p = OxmlElement('w:p')
                                            r = OxmlElement('w:r')
                                            drawing = OxmlElement('w:drawing')
                                            drawing.append(image_element)
                                            
                                            # 组装结构
                                            r.append(drawing)
                                            p.append(r)

                                            # 插入到文档结构
                                            parent.clear()
                                            parent.append(drawing)

                                        except Exception as e:
                                            print(e)
                                        
                                        os.remove(image_file_path)
                                else:
                                    element.text = element.text.replace(
                                        "{{" + key + "}}", value, 1)

                                break
                        except Exception as e:
                            element.text = ""
                            print(e)

                    if flag:
                        break

    ## 遍历文档中的所有表格
    for table in doc.tables:
        # 遍历表格中的每一行
        for row in table.rows:
            # 遍历行中的每一个单元格
            for cell in row.cells:
                # 遍历单元格中的每一个段落
                for paragraph in cell.paragraphs:
                    text = paragraph.text.replace('\n', '').replace(
                        '\r', '').replace('\r\n', '').strip()
                    if BaseClass().is_variable(text):
                        key = text.replace("{{", "").replace("}}", "")
                        # paragraph.text = paragraph.text.replace(text, info_json[key])

                        font_size = paragraph.runs[
                            0].font.size  # 假设所有格式相同，这里仅取第一个run的格式
                        color = paragraph.runs[0].font.color.rgb  # 保存颜色，如果有的话
                        bold = paragraph.runs[0].font.bold is not None  # 保存粗体状态
                        italic = paragraph.runs[
                            0].font.italic is not None  # 保存斜体状态

                        # 遍历段落中的每一个run（文本片段）
                        text_tmp = ""
                        for run in paragraph.runs:
                            # print(run.text)
                            text_tmp = text_tmp + run.text
                            if text_tmp == text:
                                # print(text_tmp)
                                # print(key)
                                # print(info_json[key])
                                # 如果run的文本包含占位符，则替换它
                                try:
                                    if ":image" in text_tmp:
                                        key_split = key.split(":")
                                        key = key_split[0]
                                        if len(key_split) == 3:
                                            size = key_split[2].split("*")
                                            width = float(size[0])
                                            height = float(size[1])
                                        else:
                                            width = None
                                            height = None

                                        if info_json[key] != "":
                                            for value_item in info_json[key]:
                                                file_token = value_item.get("file_token")
                                                name = value_item.get("name")
                                                extra = {
                                                    "bitablePerm": {
                                                        "tableId": table_id,
                                                        "attachments": {
                                                            field_id_map[key]: {
                                                                record_id:
                                                                [file_token]
                                                            }
                                                        }
                                                    }
                                                }
                                                attachment_resp = BaseClass(
                                                ).download_attachment(
                                                    personal_base_token, file_token,
                                                    extra)
                                                # print(attachment_resp)
                                                
                                                image_file_path = os.path.join(personal_main_path, name)
                                                with open(image_file_path,
                                                        'wb') as file:
                                                    file.write(
                                                        attachment_resp.content)
                                                file.close()

                                                try:
                                                    paragraph.add_run(
                                                    ).add_picture(
                                                        image_file_path,
                                                        width=Cm(width),
                                                        height=Cm(height))
                                                    # print(paragraph.text)
                                                except Exception as e:
                                                    paragraph.add_run(
                                                    ).add_picture(image_file_path)

                                                os.remove(image_file_path)

                                        run.text = ""

                                    elif ":barcode" in text_tmp:
                                        key_split = key.split(":")
                                        key = key_split[0]
                                        if len(key_split) == 3:
                                            size = key_split[2].split("*")
                                            width = float(size[0])
                                            height = float(size[1])
                                        else:
                                            width = None
                                            height = None

                                        barcode_file_name = generate_barcode(
                                            info_json[key], 'code128', None,
                                            personal_main_path)
                                        barcode_file_path = os.path.join(
                                            personal_main_path,
                                            barcode_file_name)

                                        try:
                                            paragraph.add_run().add_picture(
                                                barcode_file_path,
                                                width=Cm(width),
                                                height=Cm(height))
                                        except Exception as e:
                                            paragraph.add_run().add_picture(
                                                barcode_file_path)
                                        run.text = ""

                                    elif ":qrcode" in text_tmp:
                                        key_split = key.split(":")
                                        key = key_split[0]
                                        if len(key_split) == 3:
                                            size = key_split[2].split("*")
                                            width = float(size[0])
                                            height = float(size[1])
                                        else:
                                            width = None
                                            height = None

                                        qrcode_file_name = generate_qrcode(
                                            info_json[key], {},
                                            personal_main_path)
                                        qrcode_file_path = os.path.join(
                                            personal_main_path,
                                            qrcode_file_name)

                                        try:
                                            paragraph.add_run().add_picture(
                                                qrcode_file_path,
                                                width=Cm(width),
                                                height=Cm(height))
                                        except Exception as e:
                                            paragraph.add_run().add_picture(
                                                qrcode_file_path)
                                        run.text = ""

                                    else:
                                        run.text = info_json[key]
                                except Exception as e:
                                    run.text = ""

                                if font_size:  # 如果存在字体大小，设置字体大小
                                    run.font.size = Pt(font_size.pt)
                                if color:  # 如果存在颜色，则设置颜色
                                    run.font.color.rgb = RGBColor(*color)
                                if bold:  # 如果存在粗体，则设置粗体
                                    run.font.bold = bold
                                if italic:  # 如果存在斜体，则设置斜体
                                    run.font.italic = italic
                                text_tmp = ""
                                break  # 因为我们只处理第一个出现的占位符，所以找到后退出循环
                            else:
                                run.text = ""

    # 保存修改后的文档
    doc.save(target_file_path)

    # print(file_type)
    if file_type == 'pdf':
        #获取文件名称
        filename = target_file_path.split(".docx")[0]
        pdf_target_file_path = f"{filename}.pdf"
        system = platform.system()
        convert_flag = True

        # 将 docx 文档转换为 PDF，如果转换失败，将上传 docx 文件到附件字段中
        if system == 'Windows':
            try:
                import pythoncom
                pythoncom.CoInitialize()
                # convert(target_file_path, pdf_target_file_path)
                wordToPdf(target_file_path)

            except Exception as e:
                msg = "系统未安装 Office 软件或 PDF 转换失败，自动导出为 docx 格式"
                print(msg)
                convert_flag = False

                print(f"系统转换出错: {e}")

        elif system == 'Linux':
            command = [
                'soffice', '--infilter="Microsoft Word 2007-365"',
                '--convert-to', 'pdf', '--outdir', personal_main_path,
                target_file_path
            ]
            # print(command)
            # cmd = " ".join(command)
            # print(cmd)

            try:
                try:
                    subprocess.run(command, check=True)
                    # os.system(cmd)
                except Exception as e:
                    time.sleep(2)
                    subprocess.run(command, check=True)

            except Exception as e:
                msg = "系统未安装 LibreOffice 软件或 PDF 转换失败，自动导出为 docx 格式"
                print(msg)
                convert_flag = False

                print(f"系统转换出错: {e}")

        if not os.path.isfile(pdf_target_file_path):
            convert_flag = False

        if convert_flag:
            file = (open(pdf_target_file_path, 'rb'))
            req_body = {
                "file_name": file_name + ".pdf",
                "parent_type": "bitable_file",
                "parent_node": app_token,
                "size": str(os.path.getsize(pdf_target_file_path)),
                "file": file
            }
        else:
            file = (open(target_file_path, 'rb'))
            req_body = {
                "file_name": file_name + ".docx",
                "parent_type": "bitable_file",
                "parent_node": app_token,
                "size": str(os.path.getsize(target_file_path)),
                "file": file
            }

    # 如果 file_type 不为 `pdf` 时执行以下代码
    else:
        file = (open(target_file_path, 'rb'))
        req_body = {
            "file_name": file_name + ".docx",
            "parent_type": "bitable_file",
            "parent_node": app_token,
            "size": str(os.path.getsize(target_file_path)),
            "file": file
        }

    multi_form = MultipartEncoder(req_body)
    # 上传附件到多维表格空间
    response = BaseClass().upload_all(personal_base_token, multi_form)
    # print(response)
    file.close()

    record_list = []
    field_list = {}
    fields = {}
    if response.get("code") == 0:
        file_token = [response.get("data")]
        fields[file_field] = file_token
        field_list["fields"] = fields
        field_list["record_id"] = record_id
        record_list.append(field_list)

        # print(record_list)

        # 更新多维表格记录
        response = BaseClass().batch_update_record(app_token, personal_base_token,
                                                   table_id, record_list)
        # print(response)
        if response.get("code") == 0:
            msg = msg + "\r\n模板导出为附件成功"
            print(msg)

            # 附件更新成功后，将临时目录删除
            if os.path.isfile(target_file_path):
                file.close()
                try:
                    shutil.rmtree(personal_main_path)
                    pass
                except Exception as e:
                    print("删除临时目录出错：", e)

    return msg


def wordToPdf(word_file):
    '''
    将word文件转换成pdf文件
    :param word_file: word文件
    :return:
    '''
    # 获取word格式处理对象
    import comtypes.client
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False # 设置窗口不可见
    # 以Doc对象打开文件
    doc_ = word.Documents.Open(word_file)
    # 另存为pdf文件
    doc_.SaveAs(word_file.replace(".docx", ".pdf"), FileFormat=17)
    # 关闭doc对象
    doc_.Close()
    # 退出word对象
    word.Quit()



def create_image_element(r_id, width, height):

    """创建符合Open XML标准的图片元素结构"""
    # 单位转换（厘米转EMU）
    width_emu = int(Cm(width).emu)
    height_emu = int(Cm(height).emu)

    # print(width, height)
    # print(width_emu, height_emu)

    # ========== 核心结构构建 ==========
    # 1. 内联元素 (需确保文档根元素已声明wp命名空间)
    inline = OxmlElement('wp:inline')
    
    # 2. 设置间距属性
    inline.set('distT', "0")
    inline.set('distB', "0")
    inline.set('distL', "0")
    inline.set('distR', "0")


    # 3. 添加尺寸元素
    extent = OxmlElement('wp:extent')
    extent.set('cx', str(width_emu)) 
    extent.set('cy', str(height_emu))
    inline.append(extent)

    # 4. 文档属性（修正ID类型问题）
    docPr = OxmlElement('wp:docPr')
    docPr.set('id', "".join(re.findall(r'\d+', r_id)))  # 必须为数字，建议使用独立计数器
    docPr.set('name', "Inserted_Image")  # 使用描述性名称
    inline.append(docPr)

    # 5. 构建图形结构
    graphic = OxmlElement('a:graphic')
    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', 
                  'http://schemas.openxmlformats.org/drawingml/2006/picture')

    # 6. 图片定义
    pic = OxmlElement('pic:pic')

    # 7. 非可视化属性（修正ID设置）
    nvPicPr = OxmlElement('pic:nvPicPr')
    cNvPr = OxmlElement('pic:cNvPr')
    cNvPr.set('id', "".join(re.findall(r'\d+', r_id)))  # 不同元素使用不同ID
    cNvPr.set('name', r_id)  # 留空避免冲突
    nvPicPr.append(cNvPr)
    nvPicPr.append(OxmlElement('pic:cNvPicPr'))
    pic.append(nvPicPr)

    # 8. 图片填充设置
    blipFill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), r_id)  # 关联正确的关系ID
    blipFill.append(blip)
    
    # 添加拉伸设置
    stretch = OxmlElement('a:stretch')
    stretch.append(OxmlElement('a:fillRect'))
    blipFill.append(stretch)
    pic.append(blipFill)

    # 9. 形状属性（优化坐标设置）
    spPr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    
    # 偏移量设置
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    
    # 扩展量设置
    ext = OxmlElement('a:ext')
    ext.set('cx', str(width_emu))
    ext.set('cy', str(height_emu))
    xfrm.append(ext)
    
    spPr.append(xfrm)
    
    # 几何形状设置（关键修复点）
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')  # 必须设置prst属性
    spPr.append(prstGeom)
    
    pic.append(spPr)

    # ========== 结构组装 ==========
    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)

    return inline



## 判断文件名是否在允许的格式范围内
def allowed_file(filename):
    """
        检验文件名尾缀是否满足格式要求\r\n
        :param filename:\r\n
        :return:\r\n
    """
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


## 删除指定路径下面的所有文件
def delete_files_in_directory(directory):
    try:
        file_list = os.listdir(directory)
        try:
            file_list.remove(".gitkeep")
        except Exception as e:
            pass
        # print(file_list)
        for file_name in file_list:
            file_path = os.path.join(directory, file_name)
            if os.path.isfile(file_path):
                os.remove(file_path)
    except Exception as e:
        print(e)



## 删除模板文件接口
@app.route('/delete_templates', methods=['POST'])
def delete_file():
    """
    删除 template_files 文件夹下对应的 personal_base_token 下的所有模板文件
    """

    result_msg = "模板文件及缓存已经删除"
    
    saved_info = dict(request.form.lists()).get("saved_info")[0]
    saved_info_split = saved_info.split("_")
    template_file_path = os.path.join(app.config['UPLOAD_FOLDER'], saved_info_split[0], saved_info_split[2], saved_info_split[3])
    # print(template_file_path)

    try:
        delete_files_in_directory(template_file_path)
        os.rmdir(template_file_path)
    except Exception as e:
        print(e)
        pass


    return result_msg


## 上传模板文件接口
@app.route('/upload_template', methods=['POST'])
def upload_file():
    """
    上传文件到 template_files 文件夹下对应的 personal_base_token 下
    """

    if 'filePicker' not in request.files:
        return "不存在文件组件"

    file_list = dict(request.files.lists()).get("filePicker")

    # print(file_list)

    # print(dict(request.form.lists()))
    personal_base_token = dict(request.form.lists()).get("personal_base_token")[0]
    app_token = dict(request.form.lists()).get("app_token")[0]
    table_id = dict(request.form.lists()).get("table_id")[0]
    template_file_type = dict(request.form.lists()).get("template_file_type")[0]
    file_type = dict(request.form.lists()).get("file_type")[0]
    file_field = dict(request.form.lists()).get("file_field")[0]

    # print(personal_base_token)
    if personal_base_token == "":
        return "多维表格授权码不能为空"

    result_msg = ""

    # 每个多维表格主路径
    main_path = os.path.join(app.config['UPLOAD_FOLDER'], personal_base_token)
    # print(str(main_path))
    # print("*"*30)
    if not os.path.exists(main_path):
        os.mkdir(main_path)

    # 每个多维表格数据表主路径
    table_main_path = os.path.join(main_path, table_id)
    if not os.path.exists(table_main_path):
        os.mkdir(table_main_path)

    # 模板文件主路径
    template_file_path = os.path.join(table_main_path, template_file_type)
    # print(template_file_path)
    if not os.path.exists(template_file_path):
        os.mkdir(template_file_path)
    else:
        delete_files_in_directory(template_file_path)

    for file in file_list:
        if file.filename == '':
            return '没有选择模板文件'

        if file and allowed_file(file.filename):
            filename = "template.docx"
            try:
                file.save(os.path.join(template_file_path, filename))
                result_msg = result_msg + 'template file uploaded successfully<br><br>'
                server_url = request.headers.get("Origin")
                identifier = str(uuid.uuid1())
                # print(identifier)

                url = server_url + "/generate_attachment?identifier=" + identifier
                body = {
                    "personal_base_token": personal_base_token,
                    "app_token": app_token,
                    "table_id": table_id,
                    "file_field": file_field,
                    "file_type": file_type,
                    "template_file_type": template_file_type,
                    "record_id": "请在自动化流程中配置第一步触发记录的记录ID",
                    "file_name": "请在自动化流程配置第一步触发记录的唯一名称字段"
                }
                result_msg = result_msg + "请求URL:<br>" + url + "<br><br>请求体：<br>" + json.dumps(body, ensure_ascii=False)
            except Exception as e:
                # print(e)
                result_msg = result_msg + 'template file uploaded Fail<br>'
        else:
            result_msg = '模板文件格式不正确，请选择 docx 格式的文件'

    return result_msg


## 多维表格附件生成接口
@app.route("/generate_attachment", methods=['POST'])
def generate_attachment():

    result_msg = ""
    result_code = 200
    record_ids = []

    try:
        request_body = json.loads(request.data.decode("utf-8"))
        # print(request_body)

    except Exception as e:
        return {"msg": -1, "code": "请求参数错误"}

    
    personal_base_token = request_body.get("personal_base_token", None)
    app_token = request_body.get("app_token", None)
    table_id = request_body.get("table_id", None)
    record_id = request_body.get("record_id", None)
    file_name = request_body.get("file_name", None)
    file_field = request_body.get("file_field", None)
    file_type = request_body.get("file_type", None)
    template_file_type = request_body.get("template_file_type", "default")

    if app_token is None or personal_base_token is None or table_id is None or record_id is None or file_field is None or file_name is None:
        return {"msg": -1, "code": "请求参数为空"}

    record_ids.append(record_id)

    response = BaseClass().list_fields(app_token, personal_base_token,
                                       table_id)
    # print(response)
    field_map = {}
    field_id_map = {}
    if response.get("code") == 0:
        field_items = response.get("data").get("items")
        for item in field_items:
            field_map[item.get("field_name")] = item.get("type")
            field_id_map[item.get("field_name")] = item.get("field_id")

    # print(field_map)
    # print(field_id_map)
    # print("*" * 50)

    response = BaseClass().batch_get_records(app_token, personal_base_token,
                                             table_id, record_ids)
    # print(response)
    field_list = {}
    if response.get("code") == 0:
        records = response.get("data").get("records")[0].get("fields")
        shared_url = response.get("data").get("records")[0].get("shared_url")
        field_list["记录链接"] = shared_url
        # print(records)
        # print("*" * 50)
        for key, item in records.items():
            if key != file_field:
                field_value = BaseClass().get_field_value(field_map[key], item)
            else:
                field_value = ""

            # print(key, ":", field_value)

            field_list[key] = field_value

        try:
            msg = export_to_doc(app_token, personal_base_token, table_id,
                                record_id, field_list, file_name, file_field,
                                field_id_map, file_type, template_file_type)
            # result_msg = "生成附件成功"
            result_msg = msg

        except Exception as e:
            result_msg = "生成附件失败，请联系管理员查看日志"

    else:
        result_msg = "获取记录失败，请重试！"
        result_code = -1

    return {"msg": result_msg, "code": result_code}


## 删除 template_files 目录下面生成的条形码和二维码文件接口
@app.route("/clean_generate_files")
def clean_generate_files():
    file_path = app.config['GENERATE_FOLDER']
    delete_files_in_directory(file_path)

    timestamp = time.time()
    local_time = time.localtime(timestamp)
    formatted_local_time = time.strftime('%Y-%m-%d %H:%M:%S', local_time)

    print("【{}】template_files 目录下文件删除成功".format(formatted_local_time))
    return {"code": 200, "msg": "文件删除成功"}


## 条形码和二维码下载接口，返回文件的二进制流
@app.route("/download_file")
def download_file():
    """
    下载template_files目录下面的文件
    params:\r\n
    - file_name: 指定要下载的文件名\r\n
    - return_type: 指定返回的类型，不指定此参数默认为文件的二进制信息，可设置为 base64 生成图片的 base64 编码\r\n
    :return:
    """

    return_type = request.args.get("return_type", "file")

    file_name = request.args.get('file_name')
    file_path = os.path.join(fp, app.config['GENERATE_FOLDER'], file_name)
    if os.path.isfile(file_path):
        if return_type == 'file':
            return send_file(file_path, as_attachment=True)
        elif return_type == 'base64':
            with open(file_path, 'rb') as image_file:
                encoded_string = base64.b64encode(image_file.read())
            image_file.close()
            encoded_str = encoded_string.decode('utf-8')
            # print("Encoded Image:", encoded_str)
            return {
                "code": 200,
                "msg": "下载成功",
                "data": "data:image/png;base64," + encoded_str
            }
    else:
        return {"code": -1, "msg": "下载的文件不存在，请尝试重新生成"}


## 生成条形码接口，返回下载链接
@app.route('/generate_barcode', methods=['POST'])
def barcode():

    # "ean8": EAN8,
    # "ean8-guard": EAN8_GUARD,
    # "ean13": EAN13,
    # "ean13-guard": EAN13_GUARD,
    # "ean": EAN13,
    # "gtin": EAN14,
    # "ean14": EAN14,
    # "jan": JAN,
    # "upc": UPCA,
    # "upca": UPCA,
    # "isbn": ISBN13,
    # "isbn13": ISBN13,
    # "gs1": ISBN13,
    # "isbn10": ISBN10,
    # "issn": ISSN,
    # "code39": Code39,
    # "pzn": PZN,
    # "code128": Code128,
    # "itf": ITF,
    # "gs1_128": Gs1_128,
    # "codabar": CODABAR,
    # "nw-7": CODABAR,
    '''
    # POST 请求体参数格式如下：
    {
        "barcode_class": "code128",
        "options": {
            "module_width": 0.3,
            "module_height": 15.0,
            "font_size": 10,
            "text_distance": 5.0,
            "quiet_zone": 10.5
        }
    }
    '''

    request_body = {}

    try:
        request_body = json.loads(request.data.decode("utf-8"))

    except Exception as e:
        return '请求参数格式错误'

    # print(request_body)

    barcode_class = request_body.get("barcode_class", "code128")

    options = request_body.get("options", {})

    content = request.args.get("content", None)
    if content is None:
        content = request_body.get("content", None)
        if content is None:
            return "条码内容为空，请添加查询参数 content"

    file_path = os.path.join(fp, app.config['GENERATE_FOLDER'])

    result = generate_barcode(content, barcode_class, options, file_path)

    server_url = request.headers.get("Host")

    return {
        "code": 200,
        "msg": "生成成功",
        "data": 'https://' + server_url + '/download_file?file_name=' + result
    }


## 生成二维码接口，返回下载链接
@app.route('/generate_qrcode', methods=['POST'])
def qrcode():
    '''
    # POST 请求体参数格式如下：
    {
        "version": 2,
        "error_correction": "ERROR_CORRECT_H",
        "box_size": 12,
        "border": 2,
        "fill_color": "green",
        "back_color":"white"
    }
    '''

    request_body = {}
    try:
        request_body = json.loads(request.data.decode("utf-8"))

    except Exception as e:
        return '请求参数格式错误'

    # print(request_body)

    content = request.args.get("content", None)
    if content is None:
        content = request_body.get("content", None)
        if content is None:
            return "二维码内容为空，请添加查询参数 content"

    file_path = os.path.join(fp, app.config['GENERATE_FOLDER'])

    result = generate_qrcode(content, request_body, file_path)

    server_url = request.headers.get("Host")

    return {
        "code": 200,
        "msg": "生成成功",
        "data": 'https://' + server_url + '/download_file?file_name=' + result
    }


## 插件主页，用于上传模板文件
@app.route('/', methods=['GET'])
def index():
    identifier = str(uuid.uuid1())
    return render_template("index.html", identifier=identifier)


# 创建一个调度器
scheduler = BackgroundScheduler()
# 启动调度器
scheduler.start()
# 添加定时任务
task = scheduler.add_job(clean_generate_files,
                         'cron',
                         hour=0,
                         minute=30,
                         id='task')
# task = scheduler.add_job(clean_generate_files, 'cron', minute='*/1', id='task')

if __name__ == "__main__":
    system = platform.system()
    print('Current OS is', system)
    if system == 'Linux':
        try:
            os.system('chmod +x ./fonts_install.sh')
            os.system('./fonts_install.sh')

        except Exception as e:
            pass

    app.run(host='0.0.0.0', port=3300, debug=True, use_reloader=True)
